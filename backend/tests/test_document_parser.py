import tempfile
import unittest
from pathlib import Path

from core.document_parser import (
    CHUNK_METADATA_VERSION,
    _is_placeholder_only_section,
    _parse_docx,
    parse_markdown_content,
)
from core.document_content import normalize_document_markdown


class MarkdownParserShortContentTests(unittest.TestCase):
    def test_normalizes_collapsed_export_blocks_without_rewriting_facts(self) -> None:
        content = (
            "注意事项 * 复制<code>application.yml</code> * 修改配置。 "
            "```shell cp -R application-prod.yml application.yml ```"
        )

        normalized = normalize_document_markdown(content)

        self.assertIn("注意事项", normalized)
        self.assertIn("- 复制`application.yml`", normalized)
        self.assertIn("- 修改配置。", normalized)
        self.assertIn("```shell\ncp -R application-prod.yml application.yml\n```", normalized)

    def test_malformed_fence_language_is_kept_as_code_content(self) -> None:
        normalized = normalize_document_markdown("```shellmv portal portal8.6.6 && unzip portal.zip```")

        self.assertIn("```\nshellmv portal portal8.6.6 && unzip portal.zip\n```", normalized)

    def test_parser_records_canonical_content_version(self) -> None:
        chunks = parse_markdown_content("# 标题\n正文", "测试文档")

        self.assertEqual(
            chunks[0]["metadata"]["content_format_version"],
            "markdown-canonical.v1",
        )

    def test_preserves_short_cjk_content(self) -> None:
        chunks = parse_markdown_content("测试", "测试文档")

        self.assertEqual(len(chunks), 1)
        self.assertIn("测试", chunks[0]["content"])

    def test_preserves_short_latin_content(self) -> None:
        chunks = parse_markdown_content("OK", "Status")

        self.assertEqual(len(chunks), 1)
        self.assertIn("OK", chunks[0]["content"])

    def test_preserves_short_meaningful_content(self) -> None:
        chunks = parse_markdown_content("同意", "审批结论")

        self.assertEqual(len(chunks), 1)
        self.assertIn("同意", chunks[0]["content"])

    def test_preserves_heading_only_document(self) -> None:
        chunks = parse_markdown_content("# 测试", "测试文档")

        self.assertEqual(len(chunks), 1)
        self.assertIn("测试", chunks[0]["content"])

    def test_blank_content_remains_empty(self) -> None:
        self.assertEqual(parse_markdown_content("  \n\n\t", "空文档"), [])


class MarkdownParserPlaceholderTests(unittest.TestCase):
    def test_recognizes_exact_placeholder_values(self) -> None:
        for value in ("无", "暂无", "无内容", "不涉及", "N/A", "n/a", "-"):
            with self.subTest(value=value):
                self.assertTrue(_is_placeholder_only_section(value))

    def test_recognizes_blockquoted_template_placeholder(self) -> None:
        self.assertTrue(_is_placeholder_only_section("> 无"))
        self.assertTrue(_is_placeholder_only_section("> 暂无\n> 不涉及"))

    def test_does_not_match_meaningful_short_or_containing_text(self) -> None:
        for value in ("测试", "同意", "OK", "没有权限", "暂无处理方案"):
            with self.subTest(value=value):
                self.assertFalse(_is_placeholder_only_section(value))

    def test_filters_placeholder_only_structured_sections(self) -> None:
        content = """# 一、基本信息
所属产品：云枢

# 二、问题描述
> 无

# 三、原因分析
暂无

# 四、解决方案
开启统一失败提示。
"""

        chunks = parse_markdown_content(content, "配置说明")

        combined = "\n".join(chunk["content"] for chunk in chunks)
        self.assertIn("所属产品：云枢", combined)
        self.assertIn("开启统一失败提示", combined)
        self.assertNotIn("二、问题描述", combined)
        self.assertNotIn("三、原因分析", combined)

    def test_placeholder_only_document_does_not_use_raw_fallback(self) -> None:
        content = "# 二、问题描述\n> 无\n\n# 三、原因分析\nN/A"

        self.assertEqual(parse_markdown_content(content, "空模板"), [])

    def test_preserves_fenced_and_indented_code(self) -> None:
        for content in ("```text\nN/A\n```", "    N/A"):
            with self.subTest(content=content):
                chunks = parse_markdown_content(content, "代码示例")
                self.assertEqual(len(chunks), 1)
                self.assertIn("N/A", chunks[0]["content"])

    def test_preserves_markdown_table_with_placeholder_cell(self) -> None:
        content = """# 状态表
| 字段 | 值 |
| --- | --- |
| 说明 | 无 |
"""

        chunks = parse_markdown_content(content, "表格说明")

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0]["metadata"]["type"], "table")
        self.assertIn("| 说明 | 无 |", chunks[0]["content"])


class StructuredChunkMetadataTests(unittest.TestCase):
    def test_markdown_text_chunks_have_stable_section_metadata(self) -> None:
        content = """# 出差制度
适用于公司全体员工。

## 报销要求
出差结束后五个工作日内提交报销。
"""

        first = parse_markdown_content(content, "公司制度.md")
        second = parse_markdown_content(content, "公司制度.md")

        self.assertEqual(len(first), 2)
        self.assertEqual(
            [item["metadata"]["section_key"] for item in first],
            [item["metadata"]["section_key"] for item in second],
        )
        self.assertNotEqual(
            first[0]["metadata"]["section_key"],
            first[1]["metadata"]["section_key"],
        )
        self.assertEqual(
            first[0]["metadata"]["section_path"],
            ["公司制度.md", "出差制度"],
        )
        self.assertEqual(
            first[1]["metadata"]["section_path"],
            ["公司制度.md", "出差制度", "报销要求"],
        )
        for item in first:
            metadata = item["metadata"]
            self.assertEqual(metadata["metadata_version"], CHUNK_METADATA_VERSION)
            self.assertEqual(metadata["section_chunk_index"], 0)
            self.assertEqual(metadata["block_type"], "text")
            self.assertIn("heading", metadata)

    def test_long_markdown_table_parts_share_stable_id_and_row_ranges(self) -> None:
        row_count = 42
        rows = [
            f"| D级 | 城市{i} | 第{i}行标准说明{'费用与补贴要求' * 8} |"
            for i in range(1, row_count + 1)
        ]
        content = "\n".join([
            "# 出差标准",
            "普通员工适用以下标准。",
            "",
            "| 职级 | 城市 | 标准 |",
            "| --- | --- | --- |",
            *rows,
        ])

        chunks = parse_markdown_content(content, "公司出差管理标准.md")
        repeated = parse_markdown_content(content, "公司出差管理标准.md")
        table_chunks = [
            item for item in chunks if item["metadata"]["block_type"] == "table"
        ]
        repeated_tables = [
            item for item in repeated if item["metadata"]["block_type"] == "table"
        ]

        self.assertGreater(len(table_chunks), 1)
        table_ids = {item["metadata"]["table_id"] for item in table_chunks}
        self.assertEqual(len(table_ids), 1)
        self.assertEqual(
            table_ids,
            {item["metadata"]["table_id"] for item in repeated_tables},
        )
        self.assertEqual(
            [item["metadata"]["table_part_index"] for item in table_chunks],
            list(range(len(table_chunks))),
        )
        self.assertTrue(all(
            item["metadata"]["table_part_count"] == len(table_chunks)
            for item in table_chunks
        ))
        ranges = [
            (
                item["metadata"]["table_row_start"],
                item["metadata"]["table_row_end"],
            )
            for item in table_chunks
        ]
        self.assertEqual(ranges[0][0], 1)
        self.assertEqual(ranges[-1][1], row_count)
        for previous, current in zip(ranges, ranges[1:]):
            self.assertEqual(current[0], previous[1] + 1)

        section_chunks = [
            item
            for item in chunks
            if item["metadata"]["section_key"]
            == table_chunks[0]["metadata"]["section_key"]
        ]
        self.assertEqual(
            [item["metadata"]["section_chunk_index"] for item in section_chunks],
            list(range(len(section_chunks))),
        )
        self.assertTrue(all(
            item["metadata"]["type"] == "table" for item in table_chunks
        ))

    def test_two_tables_in_same_section_have_different_stable_ids(self) -> None:
        content = """# 标准
| 项目 | 值 |
| --- | --- |
| 住宿 | 450 |

说明文字。

| 项目 | 值 |
| --- | --- |
| 补贴 | 100 |
"""

        chunks = parse_markdown_content(content, "制度.md")
        repeated = parse_markdown_content(content, "制度.md")
        tables = [
            item for item in chunks if item["metadata"].get("type") == "table"
        ]
        repeated_tables = [
            item for item in repeated if item["metadata"].get("type") == "table"
        ]

        self.assertEqual(len(tables), 2)
        self.assertNotEqual(
            tables[0]["metadata"]["table_id"],
            tables[1]["metadata"]["table_id"],
        )
        self.assertEqual(
            [item["metadata"]["table_id"] for item in tables],
            [item["metadata"]["table_id"] for item in repeated_tables],
        )

    def test_long_docx_table_parts_share_id_and_contiguous_indexes(self) -> None:
        from docx import Document

        row_count = 38
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "travel.docx"
            document = Document()
            document.add_heading("住宿与补贴", level=1)
            document.add_paragraph("普通员工属于D级。")
            table = document.add_table(rows=1, cols=3)
            for cell, value in zip(table.rows[0].cells, ("职级", "城市", "标准")):
                cell.text = value
            for index in range(1, row_count + 1):
                cells = table.add_row().cells
                cells[0].text = "D级"
                cells[1].text = f"城市{index}"
                cells[2].text = f"住宿及补贴标准{index}{'详细报销要求' * 8}"
            document.save(path)

            chunks = _parse_docx(str(path), "公司出差标准.docx")
            repeated = _parse_docx(str(path), "公司出差标准.docx")

        tables = [
            item for item in chunks if item["metadata"]["block_type"] == "table"
        ]
        repeated_tables = [
            item for item in repeated if item["metadata"]["block_type"] == "table"
        ]
        self.assertGreater(len(tables), 1)
        self.assertEqual(len({item["metadata"]["table_id"] for item in tables}), 1)
        self.assertEqual(
            [item["metadata"]["table_id"] for item in tables],
            [item["metadata"]["table_id"] for item in repeated_tables],
        )
        self.assertEqual(
            [item["metadata"]["table_part_index"] for item in tables],
            list(range(len(tables))),
        )
        self.assertTrue(all(
            item["metadata"]["table_part_count"] == len(tables)
            for item in tables
        ))
        self.assertEqual(tables[0]["metadata"]["table_row_start"], 1)
        self.assertEqual(tables[-1]["metadata"]["table_row_end"], row_count)
        for previous, current in zip(tables, tables[1:]):
            self.assertEqual(
                current["metadata"]["table_row_start"],
                previous["metadata"]["table_row_end"] + 1,
            )
        self.assertTrue(all(
            item["metadata"]["section_path"]
            == ["公司出差标准", "住宿与补贴"]
            for item in tables
        ))


if __name__ == "__main__":
    unittest.main()
